from google.auth import impersonated_credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from docx import Document
import base64

SCOPES = ['https://www.googleapis.com/auth/gmail.readonly']
SERVICE_ACCOUNT_EMAIL = 'slwardogmp-test-sa@slwardo.iam.gserviceaccount.com'  # Replace with your service account's email

def search_gmail_and_create_doc(keyword1, keyword2, start_date, target_user_email, output_filename):
    """
    Searches Gmail for emails matching two keywords after a specified date,
    and creates a Word document with the email content.

    Args:
        keyword1 (str): The first keyword to search for.
        keyword2 (str): The second keyword to search for.
        start_date (str): The date to search from (YYYY/MM/DD).
        target_user_email (str): The email address of the user to impersonate.
        output_filename (str): The name of the Word document to create.
    """

    try:
        # Create credentials for impersonation
        creds = impersonated_credentials.Credentials(
            source_credentials=None,  # No source credentials needed for impersonation
            target_principal=SERVICE_ACCOUNT_EMAIL,
            target_scopes=SCOPES,
            subject=target_user_email,
        )

        # Refresh credentials
        creds.refresh(Request())

        # Build Gmail service
        service = build('gmail', 'v1', credentials=creds)

        # Search Gmail
        query = f'{keyword1} {keyword2} after:{start_date}'
        results = service.users().messages().list(userId='me', q=query).execute()
        messages = results.get('messages',)

        if not messages:
            print(f'No messages found matching "{query}"')
            return

        # Create Word document
        doc = Document()

        for message in messages:
            msg = service.users().messages().get(userId='me', id=message['id'], format='full').execute()
            headers = msg['payload'].get('headers',)
            subject = next((header['value'] for header in headers if header['name'] == 'Subject'), 'No Subject')
            sender = next((header['value'] for header in headers if header['name'] == 'From'), 'No Sender')
            date = next((header['value'] for header in headers if header['name'] == 'Date'), 'No Date')
            body_parts = [part for part in msg['payload'].get('parts',) if part['mimeType'] == 'text/plain']
            body = body_parts[0]['body']['data'] if body_parts else 'No body'
            body = base64.urlsafe_b64decode(body).decode('utf-8')

            doc.add_paragraph(f'Subject: {subject}')
            doc.add_paragraph(f'From: {sender}')
            doc.add_paragraph(f'Date: {date}')
            doc.add_paragraph(f'Body: {body}')
            doc.add_paragraph('----------------------------------------')

        # Save document
        doc.save(output_filename)
        print(f'Document created: {output_filename}')

    except HttpError as error:
        print(f'An error occurred: {error}')

# Prompt the user for input
keyword1 = input("Enter the first keyword: ")
keyword2 = input("Enter the second keyword: ")
start_date = input("Enter the start date (YYYY/MM/DD): ")
target_user_email = input("Enter the target user's email address: ")
output_filename = f'gmail_emails_{keyword1}_{keyword2}_{target_user_email.split("@")[0]}.docx'

search_gmail_and_create_doc(keyword1, keyword2, start_date, target_user_email, output_filename) 
